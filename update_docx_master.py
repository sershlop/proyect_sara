from docx import Document
from pathlib import Path

path = Path('C:/Users/sersh/OneDrive/Desktop/sara.2/SARA_Documento_Maestro_Extenso.docx')

doc = Document(path)

existing_index = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '2.1. NUEVAS FUNCIONES DE INTELIGENCIA Y MEJORAS':
        existing_index = i
        break

if existing_index is not None:
    # Remove the old block if it exists
    body = doc._body._element
    # remove up to 7 paragraphs starting from the heading
    for _ in range(7):
        if existing_index < len(doc.paragraphs):
            body.remove(doc.paragraphs[existing_index]._p)
    print('Removed existing section block starting at', existing_index)

insert_index = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '3. FLUJO DE PROCESAMIENTO COMPLETO':
        insert_index = i
        break

if insert_index is None:
    raise SystemExit('Insertion point not found')

body = doc._body._element
ref_p = doc.paragraphs[insert_index]._p

new_texts = [
    '2.1. NUEVAS FUNCIONES DE INTELIGENCIA Y MEJORAS',
    'Estas son las mejoras de inteligencia incorporadas en la versión actual:',
    '• embeddings.py: carga y precalienta el modelo semántico desde el arranque, reduciendo la latencia de la primera consulta. La función precalentar_modelo() fuerza una codificación inicial segura.',
    '• brain.py: mejor ranking de comandos con verificación de evidencia explícita y penalización de coincidencias débiles. Esto reduce falsos positivos en comandos similares y mejora el arbitraje.',
    '• sara.py: el arranque inicial ahora registra comandos del sistema y levanta el motor semántico antes de procesar entradas. Se informa claramente el estado del motor semántico.',
    '• database.py: nueva caché persistente de intenciones (cache_intenciones), vectores semánticos en tablas de comandos y conocimientos, e índice de archivos con métricas avanzadas como accesos, ultimo_acceso y score_relevancia.',
    '• Estas mejoras refuerzan la detección de intención, el scoring semántico y la capacidad de SARA para resolver comandos locales y búsquedas contextuales con mayor precisión.'
]

for text in new_texts:
    new_p = doc.add_paragraph(text)
    body.insert(body.index(ref_p), new_p._p)

print('Inserted corrected new section before paragraph', insert_index)
doc.save(path)
